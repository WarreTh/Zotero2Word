import warnings
import base64
import io
import tempfile
import sys
from pathlib import Path
from typing import List, Optional, Tuple, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from bs4.element import NavigableString, Tag as Bs4Tag
from bs4 import BeautifulSoup

import imgkit
from html2image import Html2Image
from config import CONFIG

def add_styled_heading(doc: Document, text: str, level: int):
    if level == 0:
        doc.add_heading(text, level=0)
    else:
        doc.add_heading(text, level=max(1, min(level, 3)))

def add_metadata_as_text(doc: Document, item):
    # Add item metadata as small italic text, then add source URL if present
    if not getattr(item, "DEFAULT_EMBED_METADATA", True):
        return
    fields_to_display: List[str] = []
    if item.item_type != "note":
        if item.creators:
            authors = "; ".join(
                f"{c.get('firstName','')} {c.get('lastName','')} ({c.get('creatorType', '')})".strip()
                for c in item.creators if isinstance(c, dict)
            )
            if authors:
                fields_to_display.append(f"Author(s): {authors}")
        if item.date:
            fields_to_display.append(f"Date: {item.date}")
    if item.date_added:
        formatted_date_added = item.date_added.replace("T", " ").replace("Z", "")
        fields_to_display.append(f"Added to Zotero: {formatted_date_added}")
    if item.tags:
        fields_to_display.append(f"Tags: {', '.join(item.tags)}")
    if not fields_to_display and not getattr(item, "meta", {}).get("data", {}).get("url", None):
        return
    # Add metadata and source as small italic text, one after another
    source_url = getattr(item, "meta", {}).get("data", {}).get("url", None)
    if fields_to_display or source_url:
        lines = []
        if fields_to_display:
            lines.append("; ".join(fields_to_display))
        if source_url:
            lines.append(f"Source: {source_url}")
        p = doc.add_paragraph()
        run = p.add_run("\n".join(lines))
        run.italic = True
        run.font.size = Pt(8)

def create_hyperlink(paragraph, url: str, text: Optional[str] = None):
    """Add a clickable hyperlink to a Word paragraph. Error checking included."""
    if not paragraph or not url or not isinstance(url, str):
        return
    if not text or not isinstance(text, str):
        text = url
    # Add relationship to document part and get rId
    doc_part = paragraph.part
    r_id = doc_part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    # Create hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    # Create run element
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)
    # Add text element
    from lxml import etree
    t = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', attrib={}, nsmap=None)
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = str(text)
    run.append(t)
    hyperlink.append(run)
    # Append hyperlink to paragraph
    paragraph._p.append(hyperlink)

def add_html_content_to_doc(doc: Document, html_content: str, DEFAULT_CODE_BLOCK_FONT_NAME="Courier New", DEFAULT_CODE_BLOCK_FONT_SIZE=10, DEFAULT_CODE_BLOCK_BG_COLOR="F0F0F0", DEFAULT_MAX_IMG_WIDTH=6.0, DEFAULT_DOWNLOAD_NOTE_IMAGES=False, verbose=True):
    if not isinstance(html_content, str):
        warnings.warn(f"HTML content must be a string, got {type(html_content)}. Skipping.")
        return
    soup = BeautifulSoup(html_content, 'html.parser')
    html_heading_formats = {
        'h1': {'size': Pt(16), 'bold': True},
        'h2': {'size': Pt(14), 'bold': True},
        'h3': {'size': Pt(12), 'bold': True},
        'h4': {'size': Pt(11), 'bold': True},
        'h5': {'size': Pt(10), 'bold': True, 'italic': True},
        'h6': {'size': Pt(10), 'italic': True},
    }
    def _add_paragraph_with_shading(doc_obj, text, font_name, font_size_pt, rgb_fg_color="000000", rgb_bg_color=None):
        p = doc_obj.add_paragraph()
        if rgb_bg_color:
            try:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), rgb_bg_color)
                p._p.get_or_add_pPr().append(shd)
            except Exception as e:
                warnings.warn(f"Could not apply background shading: {e}")
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.font.color.rgb = RGBColor.from_string(rgb_fg_color)
        return p
    def _process_node_recursively(node, current_doc_paragraph=None, current_list_style=None, current_list_level=0):
        if isinstance(node, NavigableString):
            text = str(node)
            if text.strip() or text:
                if current_doc_paragraph:
                    current_doc_paragraph.add_run(text)
                else:
                    doc.add_paragraph(text)
            return
        if not isinstance(node, Bs4Tag):
            return
        tag_name = node.name.lower()
        if tag_name in ['p', 'div']:
            p = doc.add_paragraph()
            for child in node.children:
                _process_node_recursively(child, p, current_list_style, current_list_level)
        elif tag_name in html_heading_formats:
            style = html_heading_formats[tag_name]
            p = doc.add_paragraph()
            run = p.add_run(node.get_text(strip=True))
            if style.get('bold'): run.bold = True
            if style.get('italic'): run.italic = True
            if style.get('size'): run.font.size = style['size']
        elif tag_name == 'ul':
            for li in node.find_all('li', recursive=False):
                _process_node_recursively(li, None, 'List Bullet', current_list_level + 1)
        elif tag_name == 'ol':
            for li in node.find_all('li', recursive=False):
                _process_node_recursively(li, None, 'List Number', current_list_level + 1)
        elif tag_name == 'li':
            p = doc.add_paragraph(style=current_list_style or 'ListContinue')
            if current_list_level > 1:
                p.paragraph_format.left_indent = Inches(0.25 * (current_list_level - 1))
            for child in node.children:
                _process_node_recursively(child, p)
        elif tag_name == 'pre':
            code_text = node.get_text()
            _add_paragraph_with_shading(doc, code_text, DEFAULT_CODE_BLOCK_FONT_NAME, DEFAULT_CODE_BLOCK_FONT_SIZE, rgb_bg_color=DEFAULT_CODE_BLOCK_BG_COLOR)
        elif tag_name == 'hr':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom_bdr = OxmlElement('w:bottom')
            bottom_bdr.set(qn('w:val'), 'single')
            bottom_bdr.set(qn('w:sz'), '6')
            bottom_bdr.set(qn('w:space'), '1')
            bottom_bdr.set(qn('w:color'), 'auto')
            pBdr.append(bottom_bdr)
            pPr.append(pBdr)
        elif tag_name == 'blockquote':
            p = doc.add_paragraph(style='Intense Quote')
            for child in node.children:
                _process_node_recursively(child, p)
        elif tag_name == 'img':
            src_raw = node.get('src', '')
            alt = node.get('alt', 'image')
            src = str(src_raw) if src_raw is not None else ""
            added_image = False
            if verbose:
                print(f"[Zotero2Word] Found <img> tag with src: {src}", file=sys.stderr)
            if src.startswith("data:image") and ";base64," in src:
                try:
                    img_type, base64_data = src.split(';base64,', 1)
                    img_bytes = base64.b64decode(base64_data)
                    img_stream = io.BytesIO(img_bytes)
                    if current_doc_paragraph:
                        current_doc_paragraph.add_run().add_picture(img_stream, width=Inches(min(3.0, DEFAULT_MAX_IMG_WIDTH)))
                    else:
                        doc.add_picture(img_stream, width=Inches(min(3.0, DEFAULT_MAX_IMG_WIDTH)))
                    added_image = True
                    if verbose:
                        print(f"[Zotero2Word] Inserted base64 image from HTML.", file=sys.stderr)
                except Exception as e:
                    warnings.warn(f"Error embedding base64 image: {e}")
                    if verbose:
                        print(f"[Zotero2Word] Failed to insert base64 image: {e}", file=sys.stderr)
            elif src.startswith("http://") or src.startswith("https://"):
                if verbose:
                    print(f"[Zotero2Word] External image URL found: {src}", file=sys.stderr)
                if DEFAULT_DOWNLOAD_NOTE_IMAGES:
                    if verbose:
                        print(f"[Zotero2Word] (Not implemented) Would download and insert: {src}", file=sys.stderr)
                    if current_doc_paragraph:
                        current_doc_paragraph.add_run(f"[Image: {alt} at {src}]")
                    else:
                        doc.add_paragraph(f"[Image: {alt} at {src}]")
                else:
                    if current_doc_paragraph:
                        current_doc_paragraph.add_run(f"[Image link: {src}] ({alt})")
                    else:
                        doc.add_paragraph(f"[Image link: {src}] ({alt})")
            elif src:
                if verbose:
                    print(f"[Zotero2Word] Local/relative image src found: {src}", file=sys.stderr)
                if current_doc_paragraph:
                    current_doc_paragraph.add_run(f"[Image src: {src}] ({alt})")
                else:
                    doc.add_paragraph(f"[Image src: {src}] ({alt})")
            if not added_image and not (src.startswith("http") and DEFAULT_DOWNLOAD_NOTE_IMAGES):
                if current_doc_paragraph:
                    current_doc_paragraph.add_run(f" [Image: {alt}] ")
        elif tag_name == 'br':
            if current_doc_paragraph:
                current_doc_paragraph.add_run().add_break(WD_BREAK.LINE)
        elif current_doc_paragraph:
            if tag_name in ['strong', 'b']:
                run = current_doc_paragraph.add_run(node.get_text())
                run.bold = True
            elif tag_name in ['em', 'i']:
                run = current_doc_paragraph.add_run(node.get_text())
                run.italic = True
            elif tag_name == 'u':
                run = current_doc_paragraph.add_run(node.get_text())
                run.underline = True
            elif tag_name == 'code':
                run = current_doc_paragraph.add_run(node.get_text())
                run.font.name = DEFAULT_CODE_BLOCK_FONT_NAME
            elif tag_name == 'a':
                href = str(node.get('href', ''))
                text = str(node.get_text(strip=True) or href)
                if href:
                    create_hyperlink(current_doc_paragraph, href, text)
            else:
                for child in node.children:
                    _process_node_recursively(child, current_doc_paragraph)
        else:
            temp_p = doc.add_paragraph()
            for child in node.children:
                _process_node_recursively(child, temp_p)
    for child_node in soup.contents:
        _process_node_recursively(child_node, None)

def add_image_attachment_to_doc(doc: Document, image_path: Path, DEFAULT_MAX_IMG_WIDTH=6.0, verbose=True):
    """
    Insert an image from a local file path into the Word document, with optional console logging for debugging.
    """
    # Error checking for image_path
    if image_path is None:
        warnings.warn("[ERROR] Image path is None. Cannot insert image.")
        if verbose:
            print("[Zotero2Word] Image path is None. Skipping image insertion.", file=sys.stderr)
        return
    if verbose:
        print(f"[Zotero2Word] Attempting to insert image: {image_path}", file=sys.stderr)
    if not image_path.exists():
        warnings.warn(f"Image file not found: {image_path}")
        if verbose:
            print(f"[Zotero2Word] Image file not found: {image_path}", file=sys.stderr)
        return
    try:
        doc.add_picture(str(image_path), width=Inches(DEFAULT_MAX_IMG_WIDTH))
        doc.add_paragraph()
        if verbose:
            print(f"[Zotero2Word] Successfully inserted image: {image_path}", file=sys.stderr)
    except Exception as e:
        warnings.warn(f"Failed to add image {image_path}: {e}")
        if verbose:
            print(f"[Zotero2Word] Failed to insert image: {image_path} | Error: {e}", file=sys.stderr)

def add_html_snapshot_to_doc(doc: Document, html_file_path: Path, DEFAULT_MAX_IMG_WIDTH=6.0, verbose=True):
    """
    Render HTML snapshot using html2image and insert screenshot into doc.
    Save screenshot to system temp dir with a readable name.
    Suppress GPUControl errors by disabling GPU in Chromium.
    """
    # Simple error checking for file existence
    if not html_file_path.exists():
        warnings.warn(f"Snapshot HTML file not found: {html_file_path}")
        if verbose:
            print(f"[Zotero2Word] HTML snapshot file not found: {html_file_path}", file=sys.stderr)
        return
    # Remove displaying the file:// snapshot URL (do not add it to the doc)
    # Use system temp dir and a readable, unique name
    temp_dir = Path(tempfile.gettempdir())
    base_name = html_file_path.stem
    screenshot_name = f"{base_name}.screenshot.png"
    output_image_path = temp_dir / screenshot_name
    # Major step: Use Html2Image to prepare for HTML snapshot (browser_args not supported)
    hti = Html2Image(output_path=str(temp_dir))
    try:
        # Major step: Render HTML to image
        hti.screenshot(html_file=str(html_file_path), save_as=screenshot_name, size=(1024, 1024))
        # Major step: Insert screenshot if it exists
        if output_image_path.exists() and output_image_path.stat().st_size > 0:
            doc.add_picture(str(output_image_path), width=Inches(DEFAULT_MAX_IMG_WIDTH))
            doc.add_paragraph()
            if verbose:
                print(f"[Zotero2Word] Inserted HTML snapshot image: {output_image_path}", file=sys.stderr)
        else:
            warnings.warn(f"Screenshot not found or empty after generation: {output_image_path}")
            if verbose:
                print(f"[Zotero2Word] Screenshot not found or empty: {output_image_path}", file=sys.stderr)
    except Exception as e:
        warnings.warn(f"html2image failed to convert {html_file_path}: {e}. Make sure dependencies are installed.")
        if verbose:
            print(f"[Zotero2Word] Failed to render HTML snapshot: {html_file_path} | Error: {e}", file=sys.stderr)
    # Do NOT delete the screenshot file after use

def set_paragraph_hr(paragraph):
    if not hasattr(paragraph, "_p"):
        warnings.warn("set_paragraph_hr: Provided object is not a valid Word paragraph.")
        return
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '12')
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), 'auto')
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)

def add_link_as_small_text(doc: Document, url: str):
    """Add a URL as small italic text to the document."""
    if not url:
        return
    p = doc.add_paragraph()
    run = p.add_run(url)
    run.italic = True
    run.font.size = Pt(8)

def is_image_file(filepath):
    """Check if a file is an image by extension."""
    # Error checking for filepath
    if not filepath:
        return False
    image_exts = [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]
    return str(filepath).lower().endswith(tuple(image_exts))
