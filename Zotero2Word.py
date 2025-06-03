#!/usr/bin/env python3
"""
Zotero2Word.py
Uses pyzotero (local=True) to convert your entire library into a
well-formatted Word document.
"""

from __future__ import annotations
import os
import re
import sys
import warnings
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Tuple, Any, Optional
import base64
import io
import urllib.parse
# We'll use requests for fetching remote images if needed, but make it optional for now.
import requests
import hashlib

from pyzotero import zotero
from bs4 import BeautifulSoup, Tag as Bs4Tag
from bs4.element import NavigableString
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tqdm import tqdm
import tempfile
import shutil
import imgkit # For HTML snapshots

from zotero_utils import ZItem, connect_local, get_attachment_path, populate_item_children, build_zotero_item_tree
from doc_builder import add_styled_heading, add_metadata_as_text, add_html_content_to_doc, add_image_attachment_to_doc, add_html_snapshot_to_doc, set_paragraph_hr, add_link_as_small_text

# Attempt to import configuration
try:
    from config import CONFIG
except ImportError:
    sys.exit("❌ Configuration file config.py not found. Please create it (see README).")

def check_dependencies():
    """Checks for required external command-line tools."""
    try:
        import imgkit
    except ImportError:
        sys.exit("❌ imgkit Python package is not installed. Install it with 'pip install imgkit'.")
    
    if not shutil.which("wkhtmltoimage"):
        sys.exit("❌ System binary 'wkhtmltoimage' is not installed or not in PATH. "
                 "Install it (e.g., 'sudo apt install wkhtmltopdf' on Debian/Ubuntu, "
                 "'brew install wkhtmltopdf' on macOS, or download from wkhtmltopdf.org).")

def add_table_of_contents(doc: Document, toc_structure):
    """Generates a manual Table of Contents with clickable links for each collection and subcollection.
    Only the first top-level section is shown in full; subsequent entries show only the previous level and current one, but indentation reflects depth.
    For each top-level, show full path once, then only previous and current for deeper levels."""
    doc.add_paragraph("Table of Contents", style="Title")
    shown_top_levels = set()
    for path_tuple, heading_id, bookmark_name, bookmark_num in toc_structure:
        level = len(path_tuple)
        indent = "    " * (level - 1)
        # Determine display name based on previous levels and shown top-levels
        if level == 1:
            display_name = path_tuple[-1]
            shown_top_levels.add(path_tuple[0])
        elif level == 2:
            # If top-level not shown yet, show full path, else only show previous/current
            if path_tuple[0] not in shown_top_levels:
                display_name = " / ".join(path_tuple)
                shown_top_levels.add(path_tuple[0])
            else:
                display_name = " / ".join(path_tuple[-2:])
        else:
            # For level >= 3, always show previous and current
            display_name = " / ".join([path_tuple[-2], path_tuple[-1]])
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12 * (level - 1))
        # Create the hyperlink XML
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        run.append(rPr)
        from lxml import etree
        t = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', attrib={}, nsmap=None)
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = f"{indent}{display_name}"
        run.append(t)
        hyperlink.append(run)
        p._p.append(hyperlink)
    doc.add_page_break()

def get_cached_image_path(url: str, cache_dir: str) -> str:
    """Return the path for a cached image file based on the URL hash."""
    if not url or not cache_dir:
        return None
    ext = os.path.splitext(url)[1].lower()
    if not ext or len(ext) > 6:
        ext = ".img"
    url_hash = hashlib.sha256(url.encode("utf-8")).hexdigest()
    filename = f"{url_hash}{ext}"
    return os.path.join(cache_dir, filename)

def add_link_as_small_text(doc, url):
    # Add a URL as small italic text to the document
    if not url:
        return
    p = doc.add_paragraph()
    run = p.add_run(url)
    run.italic = True
    run.font.size = Pt(8)

def main():
    check_dependencies()
    zot_local_api = connect_local(CONFIG)
    
    collections_tree, _ = build_zotero_item_tree(zot_local_api, CONFIG)
    
    doc = Document()
    doc.core_properties.title = "Zotero Library Export"
    
    # Build TOC structure: list of (path_tuple, heading_id, bookmark_name, bookmark_num)
    toc_structure = []
    heading_ids = {}
    for idx, path_tuple in enumerate(sorted(collections_tree.keys())):
        heading_id = f"toc_{idx+1}"
        bookmark_name = heading_id
        bookmark_num = str(idx+1)
        heading_ids[path_tuple] = (heading_id, bookmark_name, bookmark_num)
        toc_structure.append((path_tuple, heading_id, bookmark_name, bookmark_num))
    
    add_table_of_contents(doc, toc_structure)

    # Add main title for the document
    add_styled_heading(doc, "Zotero Library Export", level=0) # Uses "Title" style

    # Sort collections by path for consistent output
    # Path tuple sorting: (A,) (A,B) (A,C) (B,)
    sorted_collection_paths = sorted(collections_tree.keys())

    # Count total items for progress bar
    total_items = sum(len(collections_tree[coll]) for coll in sorted_collection_paths)
    item_counter = 0

    verbose = CONFIG.get('VERBOSE_LOGGING', True)

    with tqdm(total=total_items, desc="Processing Zotero Items", ncols=80) as pbar:
        for idx, coll_path_tuple in enumerate(sorted_collection_paths):
            items_in_collection = collections_tree[coll_path_tuple]
            if not items_in_collection:
                continue

            # Determine heading text and level for the collection
            if coll_path_tuple == ("Unfiled Items",):
                collection_heading_text = "Unfiled Items"
                collection_heading_level = 1 # Top-level for unfiled
            else:
                collection_heading_text = " / ".join(coll_path_tuple)
                collection_heading_level = len(coll_path_tuple) # Depth-based level
            
            # Add a bookmark for this heading
            heading_id, bookmark_name, bookmark_num = heading_ids[coll_path_tuple]
            p = doc.add_paragraph()
            run = p.add_run(collection_heading_text)
            run.bold = True
            run.font.size = Pt(15)  # Set divider title font size to 15
            # Insert bookmark XML at the start of the paragraph
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), bookmark_num)
            bookmark_start.set(qn('w:name'), bookmark_name)
            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), bookmark_num)
            p._p.insert(0, bookmark_start)
            p._p.append(bookmark_end)

            # Sort items within the collection (e.g., by date added, then those with notes first)
            items_with_notes = sorted(
                [item for item in items_in_collection if item.has_displayable_notes()],
                key=lambda x: x.date_added
            )
            items_without_notes = sorted(
                [item for item in items_in_collection if not item.has_displayable_notes()],
                key=lambda x: x.date_added
            )
            sorted_items_in_collection = items_with_notes + items_without_notes

            for item_obj in sorted_items_in_collection:
                # Add a horizontal rule before each item for clear separation
                hr_p_item = doc.add_paragraph()
                set_paragraph_hr(hr_p_item)

                # Add item title as heading
                item_title_level = collection_heading_level + 1
                add_styled_heading(doc, item_obj.title or "(No Title)", level=item_title_level)
                # Add item metadata
                add_metadata_as_text(doc, item_obj)
                # Add notes if present, removing all excessive whitespace and blank lines
                item_notes_html_list = item_obj.get_displayable_notes()
                if item_notes_html_list:
                    # Add Notes header
                    p_notes_header = doc.add_paragraph()
                    run_notes_header = p_notes_header.add_run("Notes:")
                    run_notes_header.bold = True
                    for note_html_content in item_notes_html_list:
                        # Remove all excessive whitespace and blank lines from notes
                        cleaned_html = re.sub(r"\s+", " ", note_html_content)
                        cleaned_html = re.sub(r"(\s*<br\s*/?>\s*)+", "<br>", cleaned_html, flags=re.IGNORECASE)
                        cleaned_html = cleaned_html.strip()
                        if cleaned_html:
                            add_html_content_to_doc(doc, cleaned_html, verbose=verbose)
                # Insert HTML snapshots as images, keep together with notes
                if CONFIG.get("ENABLE_WEBPAGES", True) and hasattr(item_obj, 'snapshots') and item_obj.snapshots:
                    p_snapshots_header = doc.add_paragraph()
                    run_snapshots_header = p_snapshots_header.add_run("Snapshots:")
                    run_snapshots_header.bold = True
                    for snapshot_path in item_obj.snapshots:
                        add_html_snapshot_to_doc(doc, snapshot_path, verbose=verbose)
                # Insert all image attachments (png, jpg, jpeg, gif, bmp, tiff) and keep them compact
                image_exts = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
                attachments = []
                if CONFIG.get("ENABLE_IMAGES", True) and hasattr(item_obj, 'meta') and 'attachments' in item_obj.meta:
                    attachments.extend(item_obj.meta['attachments'])
                for att in attachments:
                    att_path = get_attachment_path(att, CONFIG)
                    url = att.get("data", {}).get("url", None)
                    key = att.get("data", {}).get("key", "unknown")
                    url_added = False
                    # If local file exists and is an image, insert as before
                    if att_path and os.path.exists(att_path) and is_image_file(att_path):
                        try:
                            add_image_attachment_to_doc(doc, Path(att_path), verbose=verbose)
                            tqdm.write(f"Inserted image: {att_path}")
                        except Exception as e:
                            tqdm.write(f"Failed to add image {att_path}: {e}")
                        if url:
                            add_link_as_small_text(doc, url)
                    # If local file exists and is an HTML snapshot, check/generate screenshot in temp dir
                    elif att_path and os.path.exists(att_path) and att_path.lower().endswith(".html"):
                        temp_dir = tempfile.gettempdir()
                        base_name = os.path.splitext(os.path.basename(att_path))[0]
                        screenshot_name = f"z2w_{key}_{base_name}.screenshot.png"
                        screenshot_path = os.path.join(temp_dir, screenshot_name)
                        if os.path.exists(screenshot_path) and os.path.getsize(screenshot_path) > 0:
                            try:
                                add_image_attachment_to_doc(doc, Path(screenshot_path), verbose=verbose)
                                tqdm.write(f"Inserted existing screenshot: {screenshot_path}")
                            except Exception as e:
                                tqdm.write(f"Failed to add screenshot {screenshot_path}: {e}")
                        else:
                            try:
                                add_html_snapshot_to_doc(doc, Path(att_path), DEFAULT_MAX_IMG_WIDTH=CONFIG.get("MAX_IMG_WIDTH", 6.0), verbose=verbose)
                                if os.path.exists(screenshot_path) and os.path.getsize(screenshot_path) > 0:
                                    add_image_attachment_to_doc(doc, Path(screenshot_path), verbose=verbose)
                                    tqdm.write(f"Generated and inserted screenshot for: {att_path}")
                                else:
                                    tqdm.write(f"Screenshot not found or empty after generation: {screenshot_path}")
                            except Exception as e:
                                tqdm.write(f"Failed to generate screenshot for {att_path}: {e}")
                        if url:
                            add_link_as_small_text(doc, url)
                    # If only a URL is present and it looks like an image, check temp dir before downloading
                    elif att_path is None and url:
                        ext = os.path.splitext(url)[1].lower()
                        base_name = os.path.splitext(os.path.basename(url))[0]
                        temp_dir = tempfile.gettempdir()
                        img_name = f"z2w_{key}_{base_name}{ext}"
                        img_path = os.path.join(temp_dir, img_name)
                        if ext in image_exts or ext == ".webp":
                            if os.path.exists(img_path):
                                add_image_attachment_to_doc(doc, Path(img_path), verbose=verbose)
                                tqdm.write(f"Used existing downloaded image: {img_path}")
                            else:
                                try:
                                    tqdm.write(f"Downloading image from URL: {url}")
                                    response = requests.get(url, timeout=10)
                                    response.raise_for_status()
                                    with open(img_path, "wb") as f:
                                        f.write(response.content)
                                    add_image_attachment_to_doc(doc, Path(img_path), verbose=verbose)
                                    tqdm.write(f"Downloaded and inserted image: {url}")
                                except Exception as e:
                                    tqdm.write(f"Failed to download or add image from {url}: {e}")
                        if url:
                            add_link_as_small_text(doc, url)
                    # Always add the link as small italic text if a URL is present
                    if url and not url_added:
                        add_link_as_small_text(doc, url)
                    elif url:
                        add_link_as_small_text(doc, url)
                item_counter += 1
                pbar.update(1)

            # Add a thematic break (horizontal rule) after processing all items in a collection
            hr_p = doc.add_paragraph()
            set_paragraph_hr(hr_p)  # Adds a horizontal rule to the paragraph

    output_file_path = CONFIG["OUTPUT_DOCX"]
    try:
        doc.save(output_file_path)
        print(f"✅ Document successfully saved to: {output_file_path.resolve()}")
    except Exception as e:
        print(f"❌ Error saving document: {e}")
        error_file_path = output_file_path.with_suffix(".error.docx")
        try:
            doc.save(error_file_path)
            print(f"Attempted to save as {error_file_path}")
        except Exception as e_save_err:
            print(f"Could not save error fallback document: {e_save_err}")

if __name__ == "__main__":
    main()
